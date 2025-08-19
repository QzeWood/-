# Date： 2025/08/19 14:54
# Author: Mr. Q

import time
start_time = time.time()
print(start_time)
print("======================================")

# -*- coding: utf-8 -*-
# Text/Excel/Word 亂碼修復 + 翻譯 GUI（v4.1 + inline progress ASCII）
import tkinter as tk
from tkinter import filedialog, ttk
from pathlib import Path
import shutil
from typing import List, Tuple, Optional, Dict

# --- Progress hook（不開彈窗版本）---
PROGRESS_HOOK = None
def set_progress_hook(fn):
    """fn(stage, i=None, total=None, label=None)"""
    global PROGRESS_HOOK
    PROGRESS_HOOK = fn

# ---------- Optional deps ----------
try:
    import openpyxl
except Exception:
    openpyxl = None

try:
    import docx
except Exception:
    docx = None

try:
    from opencc import OpenCC
except Exception:
    OpenCC = None

try:
    from deep_translator import GoogleTranslator
except Exception:
    GoogleTranslator = None

# ---------- Supported types ----------
TEXT_EXTS = {".txt", ".csv", ".tsv", ".srt", ".ass", ".md", ".json",
             ".yaml", ".yml", ".log", ".ini", ".cfg", ".html", ".htm", ".xml"}
XLSX_EXTS = {".xlsx"}
DOCX_EXTS = {".docx"}

# ---------- UI <-> 內部值映射 ----------
def ui_src_mode(val: str) -> str:
    return {
        "自動": "auto",
        "簡體中文": "zh-simp",
        "繁體中文": "zh-trad",
        "日文": "ja",
        "英文": "en",
    }.get(val, "auto")

def ui_target_code(val: str) -> str:
    return {
        "不翻譯": "none",
        "簡體中文": "zh-CN",
        "繁體中文": "zh-TW",
        "英文": "en",
        "日文": "ja",
    }.get(val, "none")

# ---------- 檢測與評分 ----------
def cjk_ratio(s: str) -> float:
    if not s:
        return 0.0
    total = 0
    cjk = 0
    for ch in s:
        code = ord(ch)
        if ch.isprintable() and ch not in "\r\t":
            total += 1
            if 0x4E00 <= code <= 0x9FFF:
                cjk += 1
    return (cjk / total) if total else 0.0

def looks_mojibake(s: str) -> bool:
    non_ascii = sum(1 for ch in s if ord(ch) > 127)
    return (non_ascii > 0) and (cjk_ratio(s) < 0.02)

# ---------- 來源模式：候選編碼 / 逆轉對 ----------
def enc_candidates_for_mode(mode: str) -> List[str]:
    base = ["utf-8", "utf-16-le", "utf-16-be"]
    if mode == "zh-simp":
        return base + ["gb18030", "gbk"]
    if mode == "zh-trad":
        return base + ["big5"]
    if mode == "ja":
        return base + ["cp932", "shift_jis", "euc_jp"]
    if mode == "en":
        return base + ["cp1252", "latin1"]
    return base + ["gb18030", "gbk", "big5", "cp932", "shift_jis", "euc_jp", "cp1252", "latin1"]

def pairs_for_mode(mode: str) -> List[Tuple[str, str]]:
    zh_pairs = [
        ("cp437", "gbk"), ("latin1", "gbk"), ("cp1252", "gbk"),
        ("cp437", "cp936"), ("latin1", "cp936"), ("cp1252", "cp936"),
        ("cp437", "big5"), ("latin1", "big5"), ("cp1252", "big5"),
        ("latin1", "utf-8"), ("cp1252", "utf-8"),
    ]
    ja_pairs = [
        ("latin1", "cp932"), ("cp1252", "cp932"),
        ("latin1", "shift_jis"), ("cp437", "cp932"),
        ("cp1252", "shift_jis"),
    ]
    if mode == "zh-simp":
        return [("cp437", "gbk"), ("latin1", "gbk"), ("cp1252", "gbk"), ("latin1", "utf-8")]
    if mode == "zh-trad":
        return [("cp437", "big5"), ("latin1", "big5"), ("cp1252", "big5"), ("latin1", "utf-8")]
    if mode == "ja":
        return ja_pairs + [("latin1", "utf-8")]
    if mode == "en":
        return [("latin1", "utf-8"), ("cp1252", "utf-8")]
    return zh_pairs + ja_pairs + [("latin1", "utf-8"), ("cp1252", "utf-8")]

# ---------- mojibake 逆轉 ----------
def transform_string(bad: str, mode: str) -> str:
    pairs = pairs_for_mode(mode)
    try:
        primary = bad.encode("cp437", errors="ignore").decode("gbk", errors="ignore")
        best = primary or bad
    except Exception:
        best = bad
    best_score = cjk_ratio(best)
    if best != bad and not looks_mojibake(best):
        return best
    for wrong, right in pairs:
        try:
            t = bad.encode(wrong, errors="ignore").decode(right, errors="ignore")
            score = cjk_ratio(t)
            if score > best_score:
                best, best_score = t, score
        except Exception:
            continue
    return best

# ---------- 檔名安全修復 ----------
def safe_fix_stem(stem: str, mode: str) -> str:
    if not looks_mojibake(stem):
        return stem
    fixed = transform_string(stem, mode)
    improved = cjk_ratio(fixed) > cjk_ratio(stem)
    not_shrunk = len(fixed.strip()) >= max(1, int(len(stem.strip()) * 0.8))
    if improved and not looks_mojibake(fixed) and not_shrunk:
        return fixed
    return stem

def build_fixed_name(path: Path, mode: str) -> Path:
    parent, stem, suffix = path.parent, path.stem, path.suffix
    fixed_stem_out = safe_fix_stem(stem, mode)
    candidate = parent / f"{fixed_stem_out}_fixed{suffix}"
    i = 1
    while candidate.exists():
        candidate = parent / f"{fixed_stem_out}_fixed_{i}{suffix}"
        i += 1
    return candidate

# ---------- bytes → 最佳文本（含進度回拋） ----------
def decode_bytes_best(b: bytes, mode: str) -> Tuple[str, str]:
    global PROGRESS_HOOK
    candidates = []
    encs = enc_candidates_for_mode(mode)
    wrongs = ["latin1", "cp1252", "cp437"]
    total_steps = len(encs) + len(wrongs)
    step = 0
    if PROGRESS_HOOK:
        PROGRESS_HOOK("begin", total=total_steps, label="解碼檢測")
    for enc in encs:
        try:
            s = b.decode(enc)
            candidates.append((s, f"bytes→{enc}"))
        except Exception:
            pass
        finally:
            step += 1
            if PROGRESS_HOOK: PROGRESS_HOOK("tick", i=step, total=total_steps)
    for wrong in wrongs:
        try:
            s_bad = b.decode(wrong, errors="ignore")
            s_fix = transform_string(s_bad, mode)
            candidates.append((s_fix, f"mojibake({wrong}→*)"))
        except Exception:
            pass
        finally:
            step += 1
            if PROGRESS_HOOK: PROGRESS_HOOK("tick", i=step, total=total_steps)
    if PROGRESS_HOOK:
        PROGRESS_HOOK("end")

    if not candidates:
        try:
            s = b.decode("latin1", errors="ignore")
            return s, "fallback: latin1"
        except Exception:
            return "", "fallback: <unreadable>"

    best_s, best_tag, best_score = "", "", -1.0
    for s, tag in candidates:
        score = cjk_ratio(s) - (0.02 if looks_mojibake(s) else 0.0)
        if score > best_score:
            best_s, best_tag, best_score = s, tag, score
    return best_s, best_tag

# ---------- 翻譯（加入分段進度回拋） ----------
_TRANSLATE_CACHE: Dict[Tuple[str, str], str] = {}

def translate_text(text: str, target_lang: str) -> str:
    global PROGRESS_HOOK
    if target_lang == "none" or not text:
        return text
    if OpenCC and target_lang in ("zh-CN", "zh-TW"):
        try:
            cc = OpenCC("t2s" if target_lang == "zh-CN" else "s2t")
            chunk = 4000
            segs = [text[i:i+chunk] for i in range(0, len(text), chunk)] or [""]
            total = len(segs)
            if PROGRESS_HOOK: PROGRESS_HOOK("begin", total=total, label="中文轉換")
            outs = []
            for i, seg in enumerate(segs, 1):
                outs.append(cc.convert(seg))
                if PROGRESS_HOOK: PROGRESS_HOOK("tick", i=i, total=total)
            if PROGRESS_HOOK: PROGRESS_HOOK("end")
            return "".join(outs)
        except Exception:
            pass
    if GoogleTranslator:
        try:
            gt = GoogleTranslator(source="auto", target=target_lang)
            chunk = 4000
            segs = [text[i:i+chunk] for i in range(0, len(text), chunk)] or [""]
            total = len(segs)
            if PROGRESS_HOOK: PROGRESS_HOOK("begin", total=total, label="翻譯載入")
            outs = []
            for i, seg in enumerate(segs, 1):
                key = (seg, target_lang)
                if key in _TRANSLATE_CACHE:
                    outs.append(_TRANSLATE_CACHE[key])
                else:
                    try:
                        out = gt.translate(seg)
                    except Exception:
                        out = seg
                    _TRANSLATE_CACHE[key] = out
                    outs.append(out)
                if PROGRESS_HOOK: PROGRESS_HOOK("tick", i=i, total=total)
            if PROGRESS_HOOK: PROGRESS_HOOK("end")
            return "".join(outs)
        except Exception:
            return text
    return text

# ---------- 檔案處理 ----------
def repair_text_to_new_file(src: Path, mode: str, target: str) -> Tuple[bool, str, Path]:
    out = build_fixed_name(src, mode)
    try:
        b = src.read_bytes()
    except Exception as e:
        return False, f"無法讀取：{e}", out
    fixed, tag = decode_bytes_best(b, mode)
    fixed = translate_text(fixed, target)
    try:
        out.write_text(fixed, encoding="utf-8", errors="ignore")
        return True, tag, out
    except Exception as e:
        return False, f"寫入失敗：{e}", out

def repair_xlsx_to_new_file(src: Path, mode: str, target: str) -> Tuple[bool, str, Path]:
    out = build_fixed_name(src, mode)
    if openpyxl is None:
        return False, "未安裝 openpyxl", out
    try:
        wb = openpyxl.load_workbook(src)
    except Exception as e:
        return False, f"無法開啟：{e}", out
    changed = False
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=False):
            for cell in row:
                v = cell.value
                if isinstance(v, str) and v:
                    nv = transform_string(v, mode)
                    nv = translate_text(nv, target)
                    if nv != v:
                        cell.value = nv
                        changed = True
    try:
        wb.save(out)
        return True, ("FIXED" if changed else "COPY"), out
    except Exception as e:
        return False, f"輸出失敗：{e}", out

def repair_docx_to_new_file(src: Path, mode: str, target: str) -> Tuple[bool, str, Path]:
    out = build_fixed_name(src, mode)
    if docx is None:
        return False, "未安裝 python-docx", out
    try:
        document = docx.Document(str(src))
    except Exception as e:
        return False, f"無法開啟：{e}", out
    changed = False
    def fix_run(run):
        nonlocal changed
        if run.text:
            nv = transform_string(run.text, mode)
            nv = translate_text(nv, target)
            if nv != run.text:
                run.text = nv
                changed = True
    for para in document.paragraphs:
        for run in para.runs:
            fix_run(run)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        fix_run(run)
    try:
        document.save(str(out))
        return True, ("FIXED" if changed else "COPY"), out
    except Exception as e:
        return False, f"輸出失敗：{e}", out

def process_one(path: Path, mode: str, target: str) -> str:
    ext = path.suffix.lower()
    if ext in TEXT_EXTS:
        ok, info, out = repair_text_to_new_file(path, mode, target)
        return f"[OK] TEXT→{out} ({info}, tgt={target})" if ok else f"[ERROR] TEXT：{path} ({info})"
    if ext in XLSX_EXTS:
        ok, info, out = repair_xlsx_to_new_file(path, mode, target)
        return f"[OK] XLSX→{out} ({info}, tgt={target})" if ok else f"[ERROR] XLSX：{path} ({info})"
    if ext in DOCX_EXTS:
        ok, info, out = repair_docx_to_new_file(path, mode, target)
        return f"[OK] DOCX→{out} ({info}, tgt={target})" if ok else f"[ERROR] DOCX：{path} ({info})"
    out = build_fixed_name(path, mode)
    try:
        shutil.copy2(path, out)
        return f"[COPY] 不支援副檔名，複製為：{out}"
    except Exception as e:
        return f"[ERROR] 不支援副檔名且複製失敗：{path} -> {e}"

# ---------- GUI ----------
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("文本亂碼修復工具（v4.1_20250819）")
        self.geometry("760x420")
        self.minsize(560, 300)
        self.resizable(True, True)
        # 置中容器：把上方的兩個下拉與按鈕都裝進來
        controls = tk.Frame(self)
        controls.grid(row=0, column=0, columnspan=2, pady=(12, 6), sticky="n")  # 不拉伸，居中
        controls.grid_columnconfigure(0, weight=0)
        controls.grid_columnconfigure(1, weight=0)
        self._prog_started = False  # 單檔只啟動一次進度列
        self._prog_tag_name = "PROG_LINE"   # Text tag 名稱（用 tag 取代 mark）
        # 進度配重（預設只有解碼，保留 1% 給收尾寫檔）
        self._phase_weights = {"解碼檢測": 99, "中文轉換": 0, "翻譯載入": 0}

        # 來源語言
        self.src_lang_var = tk.StringVar(value="自動")
        tk.Label(controls, text="來源語言：").grid(row=0, column=0, sticky="e", padx=(0, 6))
        ttk.Combobox(controls, textvariable=self.src_lang_var,
                     values=["自動", "簡體中文", "繁體中文", "日文", "英文"],
                     state="readonly", width=28).grid(row=0, column=1, sticky="w")

        # 目標語言
        self.tgt_lang_var = tk.StringVar(value="不翻譯")
        tk.Label(controls, text="目標語言：").grid(row=1, column=0, sticky="e", padx=(0, 6))
        ttk.Combobox(controls, textvariable=self.tgt_lang_var,
                     values=["不翻譯", "簡體中文", "繁體中文", "英文", "日文"],
                     state="readonly", width=28).grid(row=1, column=1, sticky="w")

        # 按鈕
        tk.Button(controls, text="選擇文件並處理", command=self.process_files, width=24) \
            .grid(row=2, column=0, columnspan=2, pady=12)

        # 結果
        self.result = tk.Text(self, height=10)
        self.result.grid(row=1, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")
        self.result.configure(font=("Arial", 10))

        self.grid_rowconfigure(1, weight=1)
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure(1, weight=1)

        # 樣式
        self.result.tag_configure("HDR", font=("Arial", 10, "bold"))
        self.result.tag_configure("OK", foreground="#0a7f00")
        self.result.tag_configure("COPY", foreground="#000000")
        self.result.tag_configure("ERROR", foreground="#b00020")
        self.result.tag_configure("MUTED", foreground="#666666")

        # 置前
        self.lift()
        self.attributes("-topmost", True)
        self.after(200, lambda: self.attributes("-topmost", False))

        # 依賴狀態
        status = []
        status.append("OpenCC: OK" if OpenCC else "OpenCC: 未安裝")
        status.append("GoogleTranslator: OK" if GoogleTranslator else "GoogleTranslator: 未安裝/不可用")
        tk.Label(self, text=" / ".join(status), fg="#666").grid(row=2, column=0, columnspan=2, pady=(0,8))

    # ====== 文字條進度：用 Text mark 精準覆寫 ======
    def _ascii_bar(self, pct: int, side: int = 16) -> str:
        pct = max(0, min(100, int(pct)))
        filled = max(1, int(round(pct * side / 100)))
        left  = "|" * filled + " " * (side - filled)
        right = " " * (side - filled) + "|" * filled
        return f"{left}{pct:>3d}%{right}"

    def begin_progress(self, label: str = "處理中…"):
        # 先把舊的進度區間 tag 移除（不存在也沒關係）
        try:
            self.result.tag_delete(self._prog_tag_name)
        except tk.TclError:
            pass
        self._prog_label = label
        self._prog_started = True
        self._prog_line_created = False

    def update_progress(self, pct: int):
        tag = self._prog_tag_name

        # 第一次 tick：先插標題行，再插「被 tag 標記」的進度條本體 + 換行
        if not getattr(self, "_prog_line_created", False):
            self.result.insert(tk.END, f"{getattr(self, '_prog_label', '處理中…')}\n", "MUTED")
            self.result.insert(tk.END, self._ascii_bar(pct), (tag, "MUTED"))  # 只給進度條本體加 tag
            self.result.insert(tk.END, "\n")  # 換行獨立，避免被覆蓋
            self._prog_line_created = True
            self.result.see(tk.END)
            self.update_idletasks()
            return

        # 後續 tick：精準覆寫「tag 範圍」內容（單行、不新增）
        ranges = self.result.tag_ranges(tag)
        if ranges:
            start, end = ranges[0], ranges[1]
            self.result.delete(start, end)
            self.result.insert(start, self._ascii_bar(pct), (tag, "MUTED"))
            self.result.see(tk.END)
            self.update_idletasks()

    def end_progress(self, msg: str = "完成"):
        self.update_progress(100)
        self.result.insert(tk.END, f"  {msg}\n", "MUTED")
        # 立刻把 100% + 完成 推到 UI
        self.result.see(tk.END)
        self.update_idletasks()

        self._prog_started = False
        self._prog_line_created = False

    def process_files(self):
        mode = ui_src_mode(self.src_lang_var.get())
        target = ui_target_code(self.tgt_lang_var.get())

        files = filedialog.askopenfilenames(
            title="選擇需要修復的文件",
            filetypes=[("All Supported", "*.txt *.csv *.tsv *.srt *.ass *.md *.json *.yaml *.yml *.log *.ini *.cfg *.html *.htm *.xml *.xlsx *.docx"),
                       ("Text", "*.txt *.csv *.tsv *.srt *.ass *.md *.json *.yaml *.yml *.log *.ini *.cfg *.html *.htm *.xml"),
                       ("Excel", "*.xlsx"),
                       ("Word", "*.docx"),
                       ("All files", "*.*")]
        )
        if not files:
            return
        # 清空，確保進度條在最上方
        self.result.delete("1.0", tk.END)

        results = []
        for f in files:
            p = Path(f)
            # Hook：解碼/翻譯共用同一條
            # 檔案處理迴圈內（for f in files: 的上方/內部都可），用這個覆蓋你原本的 _hook
            self._overall_base = 0  # 已完成的子階段累積（0~99）
            self._seg_total = 1  # 當前子階段總步數
            self._seg_weight = 0  # 當前子階段配重
            self._current_seg_label = ""
            # 依當次任務設定子階段配重（總和 ≤ 99）
            if target == "none":
                # 只做解碼：把 99% 都給解碼，最後 1% 給寫出/收尾
                self._phase_weights = {"解碼檢測": 99, "中文轉換": 0, "翻譯載入": 0}
            elif target in ("zh-CN", "zh-TW") and OpenCC:
                # 解碼 + 中文內部轉換
                self._phase_weights = {"解碼檢測": 70, "中文轉換": 29, "翻譯載入": 0}
            else:
                # 解碼 + 外語翻譯
                self._phase_weights = {"解碼檢測": 60, "翻譯載入": 39, "中文轉換": 0}

            def _hook(stage, i=None, total=None, label=None, _p=p):
                if stage == "begin":
                    if not getattr(self, "_prog_started", False):
                        self.begin_progress(label or f"處理：{_p.name}")
                    self._current_seg_label = label or ""
                    self._seg_total = max(1, total or 1)
                    # 這個子階段最多只能把總條推進到 99% 之內
                    # 找不到對應標籤時，用「剩餘配重」當預設，避免卡住
                    self._seg_weight = self._phase_weights.get(
                        self._current_seg_label,
                        max(0, 99 - self._overall_base)
                    )

                elif stage == "tick":
                    # 子階段內部進度 → 映射到 [overall_base, overall_base + seg_weight]，最高封頂 99
                    pct_in_seg = int((i or 0) * 100 / max(1, self._seg_total))
                    overall = min(99, int(self._overall_base + self._seg_weight * pct_in_seg / 100))
                    self.update_progress(overall)
                elif stage == "end":
                    # 子階段結束，把 base 往前推，但依舊封頂 99（留 1% 給「寫入完成」）
                    self._overall_base = min(99, self._overall_base + self._seg_weight)

            set_progress_hook(_hook)

            try:
                msg = process_one(p, mode, target)
            except Exception as e:
                msg = f"[ERROR] 例外：{f} -> {e}"
            finally:
                set_progress_hook(None)
                if getattr(self, "_prog_started", False):
                    self.end_progress("完成")

            print(msg)
            results.append(msg)

        # --- UX 輸出 ---
        ok_list = [r for r in results if r.startswith("[OK]")]
        copy_list = [r for r in results if r.startswith("[COPY]")]
        err_list = [r for r in results if r.startswith("[ERROR]")]
        summary_line = (f"✅ 成功 {len(ok_list)}"
                        f"   | 📄 複製 {len(copy_list)}"
                        f"   | ❌ 失敗 {len(err_list)}")

        self.result.insert(tk.END, summary_line + "\n", "HDR")
        self.result.insert(tk.END, f"來源模式={mode}, 目標語言={target}\n\n", "MUTED")

        def dump_block(title_emoji, title, lines, tag):
            if not lines:
                return
            self.result.insert(tk.END, f"{title_emoji} {title}（{len(lines)}）\n", "HDR")
            for line in lines[:8]:
                self.result.insert(tk.END, "  • " + line + "\n", tag)
            if len(lines) > 8:
                self.result.insert(tk.END, f"  … 還有 {len(lines) - 8} 條\n", "MUTED")
            self.result.insert(tk.END, "\n")

        dump_block("✅", "修復完成", ok_list, "OK")
        dump_block("📄", "僅複製（無內容變更/不支援副檔名）", copy_list, "COPY")
        dump_block("❌", "處理失敗", err_list, "ERROR")

if __name__ == "__main__":
    App().mainloop()

print("======================================")
print("ok!!")
end_time = time.time()
print(end_time)
duration = end_time - start_time
print(f"程式執行時間為 {duration:.2f} 秒")
